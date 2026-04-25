"""
Generate Meta Ads Automation - Full System Documentation (Word .docx)
Run: python scripts/generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Theme colors ──────────────────────────────────────────────
BLUE      = RGBColor(0x18, 0x77, 0xF2)   # Meta blue
DARK      = RGBColor(0x1C, 0x1E, 0x21)   # Near black
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF0, 0xF2, 0xF5)
GREEN     = RGBColor(0x2E, 0x7D, 0x32)
RED       = RGBColor(0xC6, 0x28, 0x28)
ORANGE    = RGBColor(0xE6, 0x5C, 0x00)
GRAY      = RGBColor(0x60, 0x67, 0x70)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=BLUE):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(doc, text, bold=False, italic=False, color=None, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    # Light gray background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F2F5')
    pPr.append(shd)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, header_color='1877F2'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'FAFAFA')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1877F2')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p

# ════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run('\n\n')

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('META ADS AUTOMATION')
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Complete System Documentation')
r2.font.size  = Pt(16)
r2.font.color.rgb = DARK

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run('Built with Meta Marketing API + Claude AI + Python')
r3.font.size  = Pt(12)
r3.font.color.rgb = GRAY
r3.italic = True

doc.add_paragraph('\n\n')

info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'
info_data = [
    ('Version',    'v1.0'),
    ('Date',       '25 April 2026'),
    ('Platform',   'Windows / Python 3.9+'),
    ('API',        'Meta Marketing API v21.0'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    set_cell_bg(info_table.rows[i].cells[0], '1877F2')
    for para in info_table.rows[i].cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(11)
    for para in info_table.rows[i].cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, 'Table of Contents', level=1)
toc_items = [
    '1. What Is This System?',
    '2. System Architecture',
    '3. Project File Structure',
    '4. Setup & Installation',
    '5. Credentials Required',
    '6. All CLI Commands',
    '7. Module 1 - Campaign Management',
    '8. Module 2 - Budget Optimizer',
    '9. Module 3 - Performance Reports',
    '10. Module 4 - Alert System',
    '11. Module 5 - AI Ad Copy Generator',
    '12. Module 6 - Local Database',
    '13. Optimization Rules (Detailed)',
    '14. Configuration & Thresholds',
    '15. Full Day-to-Day Workflow',
    '16. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS THIS SYSTEM
# ════════════════════════════════════════════════════════════════
add_heading(doc, '1. What Is This System?', level=1)
add_divider(doc)
add_para(doc, 'Meta Ads Automation is a command-line tool built in Python that connects directly to the Meta Marketing API to automate the creation, management, monitoring, and optimization of your Facebook and Instagram advertising campaigns.', size=11)
doc.add_paragraph()
add_para(doc, 'Instead of manually logging into Meta Ads Manager every day to check performance, adjust budgets, pause underperforming ads, or generate ad copy - this tool does it all automatically from your computer terminal.', size=11)
doc.add_paragraph()
add_heading(doc, 'What Problems It Solves', level=2)
add_bullet(doc, 'Saves 1-2 hours daily of manual campaign checking')
add_bullet(doc, 'Never misses a spend spike or ROAS drop - instant alerts')
add_bullet(doc, 'Automatically scales winning ads and pauses losing ones')
add_bullet(doc, 'Generates AI-written ad copy in seconds')
add_bullet(doc, 'Keeps full history of every change ever made')
add_bullet(doc, 'Exports performance reports as tables, CSV, or JSON')
doc.add_paragraph()
add_heading(doc, 'What It Is NOT', level=2)
add_bullet(doc, 'Not a web app - it runs from your terminal/command prompt')
add_bullet(doc, 'Not a replacement for Meta Ads Manager - it works alongside it')
add_bullet(doc, 'Not a set-and-forget system - you still decide strategy and budgets')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 2 — SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '2. System Architecture', level=1)
add_divider(doc)
add_para(doc, 'The system is built in 6 interconnected layers:', size=11)
doc.add_paragraph()

add_code(doc, 'YOU  (type commands in terminal)')
add_code(doc, '  |')
add_code(doc, '  v')
add_code(doc, 'main.py  (CLI entry point - receives your commands)')
add_code(doc, '  |')
add_code(doc, '  +---> api/          (talks to Meta Marketing API)')
add_code(doc, '  |       |')
add_code(doc, '  |       v')
add_code(doc, '  |    Facebook/Instagram (your real ads on Meta)')
add_code(doc, '  |')
add_code(doc, '  +---> database/     (saves everything locally in SQLite)')
add_code(doc, '  |')
add_code(doc, '  +---> campaigns/    (campaign creation + optimization rules)')
add_code(doc, '  |')
add_code(doc, '  +---> reporting/    (pulls metrics, formats reports, exports)')
add_code(doc, '  |')
add_code(doc, '  +---> monitor/      (alert engine + scheduled checks)')
add_code(doc, '  |')
add_code(doc, '  +---> creative/     (AI-powered ad copy generation)')

doc.add_paragraph()
add_para(doc, 'Data Flow:', bold=True, size=12)
add_bullet(doc, 'You type a command  ->  main.py parses it  ->  calls the right module')
add_bullet(doc, 'Module calls Meta API  ->  gets/creates data  ->  saves to local database')
add_bullet(doc, 'Reports read from local database (fast, no API call needed)')
add_bullet(doc, 'Monitor runs continuously in background, checks every N minutes')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 3 — FILE STRUCTURE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '3. Project File Structure', level=1)
add_divider(doc)
add_para(doc, 'Location: C:\\Users\\HP\\Desktop\\Meta Ads Automation\\', italic=True, color=GRAY)
doc.add_paragraph()

add_code(doc, 'Meta Ads Automation/')
add_code(doc, '|-- main.py                   <- CLI entry point, all commands')
add_code(doc, '|-- config.py                 <- All settings, thresholds, constants')
add_code(doc, '|-- requirements.txt          <- Python packages needed')
add_code(doc, '|-- .env                      <- YOUR credentials (never share this)')
add_code(doc, '|-- .env.example              <- Template showing what .env needs')
add_code(doc, '|-- setup.bat                 <- Windows installer script')
add_code(doc, '|-- meta_ads.db               <- Local database (auto-created)')
add_code(doc, '|-- automation.log            <- Log file (auto-created)')
add_code(doc, '|')
add_code(doc, '|-- api/                      <- Meta Marketing API wrappers')
add_code(doc, '|   |-- client.py             <- API initialisation & auth')
add_code(doc, '|   |-- campaigns.py          <- Create/pause/resume campaigns')
add_code(doc, '|   |-- adsets.py             <- Create/update ad sets & budgets')
add_code(doc, '|   |-- ads.py                <- Create ad creatives & ads')
add_code(doc, '|   |-- insights.py           <- Pull performance metrics')
add_code(doc, '|')
add_code(doc, '|-- campaigns/                <- Campaign logic')
add_code(doc, '|   |-- creator.py            <- Full campaign creation flow')
add_code(doc, '|   |-- optimizer.py          <- Budget optimization rules engine')
add_code(doc, '|   |-- validator.py          <- Input validation')
add_code(doc, '|')
add_code(doc, '|-- creative/                 <- AI ad copy generation')
add_code(doc, '|   |-- copy_generator.py     <- Claude AI integration')
add_code(doc, '|   |-- prompts/              <- Prompt templates for Claude')
add_code(doc, '|')
add_code(doc, '|-- monitor/                  <- Monitoring & alerts')
add_code(doc, '|   |-- alert_engine.py       <- Threshold checks & alert triggers')
add_code(doc, '|   |-- scheduler.py          <- Continuous monitoring loop')
add_code(doc, '|')
add_code(doc, '|-- reporting/                <- Reports & exports')
add_code(doc, '|   |-- collector.py          <- Fetches metrics from Meta API')
add_code(doc, '|   |-- formatter.py          <- Formats data as colored tables')
add_code(doc, '|   |-- exporter.py           <- Exports to CSV/JSON files')
add_code(doc, '|')
add_code(doc, '|-- database/                 <- Local data storage')
add_code(doc, '|   |-- db.py                 <- All database operations')
add_code(doc, '|')
add_code(doc, '|-- utils/                    <- Shared utilities')
add_code(doc, '|   |-- logger.py             <- Color-coded logging')
add_code(doc, '|   |-- helpers.py            <- Currency formatting, date parsing')
add_code(doc, '|   |-- mailer.py             <- Email alert sender')
add_code(doc, '|')
add_code(doc, '|-- scripts/')
add_code(doc, '    |-- seed_test_data.py     <- Loads sample data for testing')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 4 — SETUP
# ════════════════════════════════════════════════════════════════
add_heading(doc, '4. Setup & Installation', level=1)
add_divider(doc)

add_heading(doc, 'Step 1 - Install Dependencies', level=2)
add_para(doc, 'Double-click setup.bat OR run in terminal:', size=11)
add_code(doc, 'pip install -r requirements.txt')
doc.add_paragraph()
add_para(doc, 'Packages installed:', bold=True, size=11)
add_table(doc,
    ['Package', 'Purpose'],
    [
        ('facebook-business', 'Official Meta Marketing API SDK'),
        ('anthropic',         'Claude AI for ad copy generation (optional)'),
        ('colorama',          'Colored terminal output'),
        ('tabulate',          'Formatted data tables in terminal'),
        ('python-dotenv',     'Loads .env credentials file'),
        ('schedule',          'Runs monitor checks on a timer'),
        ('pandas',            'Data processing for reports'),
        ('requests',          'HTTP client'),
        ('openpyxl',          'Excel export support'),
    ]
)

add_heading(doc, 'Step 2 - Create .env File', level=2)
add_code(doc, 'copy .env.example .env')
add_para(doc, 'Then open .env and fill in your credentials (see Section 5).', size=11)

add_heading(doc, 'Step 3 - Verify Connection', level=2)
add_code(doc, 'python main.py setup verify')
add_para(doc, 'If successful, you will see your account name, currency, and status printed in green.', size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 5 — CREDENTIALS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '5. Credentials Required', level=1)
add_divider(doc)
add_para(doc, 'All credentials are stored in your .env file. NEVER share this file or commit it to Git.', bold=True, color=RED, size=11)
doc.add_paragraph()

add_table(doc,
    ['Credential', 'Where to Get It', 'Required?'],
    [
        ('META_APP_ID',        'developers.facebook.com -> Your App -> Settings -> Basic', 'Yes'),
        ('META_APP_SECRET',    'developers.facebook.com -> Your App -> Settings -> Basic -> Show', 'Yes'),
        ('META_ACCESS_TOKEN',  'developers.facebook.com/tools/explorer -> Generate Token', 'Yes'),
        ('META_AD_ACCOUNT_ID', 'business.facebook.com -> Settings -> Ad Accounts (numbers only)', 'Yes'),
        ('META_PAGE_ID',       'Facebook Page -> About -> scroll to bottom -> Page ID', 'Yes'),
        ('ANTHROPIC_API_KEY',  'console.anthropic.com (only needed for AI copy feature)', 'No'),
        ('SMTP_USER/PASS',     'Your Gmail + App Password (for email alerts)', 'No'),
    ]
)

add_heading(doc, 'Access Token Important Notes', level=2)
add_bullet(doc, 'Short-lived token expires in 1 hour - use for testing only')
add_bullet(doc, 'Long-lived token lasts 60 days - go to developers.facebook.com/tools/accesstoken to extend')
add_bullet(doc, 'System User Token never expires - recommended for production use')
add_bullet(doc, 'Required permissions: ads_management, ads_read, business_management, pages_read_engagement')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 6 — ALL CLI COMMANDS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '6. All CLI Commands', level=1)
add_divider(doc)
add_para(doc, 'All commands follow this pattern: python main.py <command> <action> [options]', italic=True, size=11)
doc.add_paragraph()

add_heading(doc, 'Campaign Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py campaign create --name "..." --objective OUTCOME_TRAFFIC --budget 500', 'Creates a new campaign + ad set'),
        ('python main.py campaign create ... --dry-run', 'Previews campaign creation without spending'),
        ('python main.py campaign list', 'Lists all campaigns from local database'),
        ('python main.py campaign list --sync', 'Syncs from Meta API first, then lists'),
        ('python main.py campaign pause --id CAMPAIGN_ID', 'Pauses a campaign'),
        ('python main.py campaign resume --id CAMPAIGN_ID', 'Resumes a paused campaign'),
        ('python main.py campaign delete --id CAMPAIGN_ID --confirm', 'Permanently deletes campaign'),
    ]
)

add_heading(doc, 'Ad Set Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py adset create --campaign-id ID --name "..." --budget 300', 'Creates a new ad set'),
        ('python main.py adset list --campaign-id ID', 'Lists all ad sets for a campaign'),
        ('python main.py adset pause --id ADSET_ID', 'Pauses an ad set'),
        ('python main.py adset resume --id ADSET_ID', 'Resumes an ad set'),
    ]
)

add_heading(doc, 'Ad Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py ad create --adset-id ID --campaign-id ID --name "..." --headline "..." --body "..." --cta SHOP_NOW --link https://...', 'Creates a new ad'),
        ('python main.py ad list --adset-id ID', 'Lists all ads in an ad set'),
    ]
)

add_heading(doc, 'Optimization Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py optimize', 'Runs all optimization rules on active campaigns'),
        ('python main.py optimize --dry-run', 'Shows what WOULD be done, no changes made'),
        ('python main.py optimize --min-roas 2.0 --max-cpa 1500', 'Override default thresholds'),
    ]
)

add_heading(doc, 'Report Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py report --campaign-id ID --date-range last_7d', 'Campaign report as table'),
        ('python main.py report --campaign-id ID --format csv --output report.csv', 'Export to CSV file'),
        ('python main.py report --campaign-id ID --format json --output report.json', 'Export to JSON file'),
        ('python main.py report account --date-range last_30d', 'Account-level report across all campaigns'),
    ]
)

add_heading(doc, 'Monitor Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py monitor start --interval 60', 'Starts continuous monitoring every 60 minutes'),
        ('python main.py monitor start --alert-email you@gmail.com', 'Monitoring with email alerts'),
        ('python main.py monitor check', 'One-time alert check right now'),
    ]
)

add_heading(doc, 'Copy & Utility Commands', level=2)
add_table(doc,
    ['Command', 'What It Does'],
    [
        ('python main.py copy generate --product "..." --audience "..." --tone urgent --count 3', 'Generate 3 AI ad copy variations'),
        ('python main.py copy list', 'List all saved copy variations'),
        ('python main.py stats', 'Show database row counts for all tables'),
        ('python main.py setup verify', 'Test Meta API connection'),
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 7 — CAMPAIGN MANAGEMENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '7. Module 1 - Campaign Management', level=1)
add_divider(doc)
add_para(doc, 'This module handles creating and managing the full Meta ads structure.', size=11)
doc.add_paragraph()

add_heading(doc, 'The 3-Level Ad Structure', level=2)
add_code(doc, 'CAMPAIGN  (the overall goal)')
add_code(doc, '    |')
add_code(doc, '    +-- AD SET  (who sees it, when, how much per day)')
add_code(doc, '          |')
add_code(doc, '          +-- AD  (the actual creative: image + headline + text)')

doc.add_paragraph()
add_table(doc,
    ['Level', 'What You Set Here', 'Example'],
    [
        ('Campaign',  'Goal/Objective, overall budget', 'Get website traffic, Rs.500/day'),
        ('Ad Set',    'Target audience, schedule, optimization', 'Women 25-35, Mumbai, interested in fitness'),
        ('Ad',        'Headline, body text, image, CTA button', '"Get 50% Off Today" + shoe image + Shop Now'),
    ]
)

add_heading(doc, 'Campaign Objectives Available', level=2)
add_table(doc,
    ['Objective', 'Use When You Want...'],
    [
        ('OUTCOME_TRAFFIC',        'People to visit your website'),
        ('OUTCOME_LEADS',          'People to fill a contact form or sign up'),
        ('OUTCOME_SALES',          'People to purchase (requires Pixel setup)'),
        ('OUTCOME_AWARENESS',      'Maximum people to see your brand'),
        ('OUTCOME_ENGAGEMENT',     'Likes, comments, shares on your post'),
        ('OUTCOME_APP_PROMOTION',  'App installs or in-app actions'),
    ]
)

add_heading(doc, 'Safety Feature - Campaigns Created as PAUSED', level=2)
add_para(doc, 'Every campaign created by this tool is automatically set to PAUSED status. This is intentional - it prevents any accidental ad spend. You manually activate campaigns in Meta Ads Manager once you have reviewed everything.', size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 8 — BUDGET OPTIMIZER
# ════════════════════════════════════════════════════════════════
add_heading(doc, '8. Module 2 - Budget Optimizer', level=1)
add_divider(doc)
add_para(doc, 'This is the most powerful feature. It automatically analyzes the last 7 days of performance for every active ad set and takes action based on predefined rules.', size=11)
doc.add_paragraph()

add_heading(doc, 'How It Works', level=2)
add_code(doc, '1. Fetch all ACTIVE campaigns from database')
add_code(doc, '2. For each campaign, get all ACTIVE ad sets')
add_code(doc, '3. For each ad set, pull last 7 days of Meta Insights data')
add_code(doc, '4. Calculate: avg ROAS, avg CTR, avg frequency, total spend, CPL')
add_code(doc, '5. Run through 7 rules (in priority order)')
add_code(doc, '6. Take action (or log as no_action)')
add_code(doc, '7. Save action to optimization_log table')
add_code(doc, '8. If --dry-run, only log - do not call Meta API')

doc.add_paragraph()
add_heading(doc, 'The 7 Optimization Rules', level=2)
add_table(doc,
    ['Priority', 'Rule Name', 'Condition', 'Action Taken'],
    [
        ('1', 'ROAS Guard',         'ROAS < 1.5 for 3+ days AND spent Rs.4000+',            'PAUSE the ad set'),
        ('2', 'CPA Guard',          'Cost per lead > Rs.2000 AND 3+ leads collected',        'Decrease budget by 20%'),
        ('3', 'High Performer',     'ROAS >= 3.0 AND cost per lead <= Rs.1250',              'Increase budget by 15%'),
        ('4', 'Low CTR Alert',      'CTR < 0.5% AND 5000+ impressions shown',               'Alert (refresh creative)'),
        ('5', 'Frequency Fatigue',  'Same people seeing ad 4+ times on average',            'Alert (audience tired)'),
        ('6', 'Budget Depleted',    '95% of budget spent before 6 PM',                      'Critical alert + email'),
        ('7', 'Spend Spike',        "Today's spend is 2x higher than 7-day daily average",  'Critical alert + email'),
    ]
)

doc.add_paragraph()
add_heading(doc, 'Budget Change Limits (Safety Guards)', level=2)
add_bullet(doc, 'Minimum ad set budget: Rs.40 per day (never goes below this)')
add_bullet(doc, 'Maximum budget increase: 5x the original budget (cannot scale infinitely)')
add_bullet(doc, 'Budget decrease: always 20% at a time (gradual, not sudden)')
add_bullet(doc, 'ROAS must be below threshold for 3 consecutive days before pausing (not just one bad day)')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 9 — REPORTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '9. Module 3 - Performance Reports', level=1)
add_divider(doc)

add_heading(doc, 'Metrics Tracked', level=2)
add_table(doc,
    ['Metric', 'What It Means', 'Good Value'],
    [
        ('Impressions',       'How many times your ad was shown',              'Depends on budget'),
        ('Reach',             'How many unique people saw your ad',            'Close to impressions'),
        ('Clicks',            'How many people clicked your ad',               'More is better'),
        ('Spend',             'How much money was spent (in your currency)',   'Within your budget'),
        ('CTR',               'Click-Through Rate (clicks / impressions x100)', '> 1.5% is good'),
        ('CPC',               'Cost Per Click',                                'Lower is better'),
        ('CPM',               'Cost per 1000 impressions',                     'Lower is better'),
        ('ROAS',              'Return on Ad Spend (revenue / spend)',          '> 2.0 is good'),
        ('Frequency',         'Average times one person saw your ad',          '< 4.0 is safe'),
        ('Leads',             'Number of form submissions / sign-ups',         'More is better'),
        ('Purchases',         'Number of completed purchases',                 'More is better'),
        ('Cost per Lead',     'How much each lead costs you',                  'Lower is better'),
        ('Cost per Purchase', 'How much each purchase costs you',              'Lower than profit margin'),
    ]
)

add_heading(doc, 'Color Coding in Table Reports', level=2)
add_table(doc,
    ['Color', 'Meaning', 'Example Thresholds'],
    [
        ('GREEN',  'Performing well',      'ROAS > 2.0, CTR > 1.5%'),
        ('YELLOW', 'Borderline - watch it', 'ROAS 1.0-2.0, CTR 0.5-1.5%'),
        ('RED',    'Underperforming',       'ROAS < 1.0, CTR < 0.5%'),
    ]
)

add_heading(doc, 'Date Range Options', level=2)
add_code(doc, 'today          - just today')
add_code(doc, 'yesterday      - just yesterday')
add_code(doc, 'last_7d        - last 7 days')
add_code(doc, 'last_30d       - last 30 days')
add_code(doc, 'last_90d       - last 90 days')
add_code(doc, 'this_month     - current month so far')
add_code(doc, 'last_month     - full previous month')
add_code(doc, '2025-04-01:2025-04-30  - custom date range (YYYY-MM-DD:YYYY-MM-DD)')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 10 — ALERTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '10. Module 4 - Alert System', level=1)
add_divider(doc)

add_heading(doc, 'Alert Severity Levels', level=2)
add_table(doc,
    ['Level', 'Color', 'Triggers Email?', 'Alert Types'],
    [
        ('CRITICAL', 'Red',    'Yes (if SMTP configured)', 'spend_spike, budget_depleted, roas < 0.5'),
        ('WARNING',  'Yellow', 'No (console only)',        'high_cpa, low_roas, low_ctr'),
        ('INFO',     'Cyan',   'No (console only)',        'frequency_fatigue'),
    ]
)

add_heading(doc, 'Email Alert Setup (Optional)', level=2)
add_para(doc, 'To receive email alerts, add these to your .env file:', size=11)
add_code(doc, 'SMTP_HOST=smtp.gmail.com')
add_code(doc, 'SMTP_PORT=587')
add_code(doc, 'SMTP_USER=your@gmail.com')
add_code(doc, 'SMTP_PASS=your_gmail_app_password')
add_code(doc, 'ALERT_EMAIL=where_alerts_go@email.com')
doc.add_paragraph()
add_para(doc, 'For Gmail: enable 2-Step Verification, then generate an App Password at myaccount.google.com/apppasswords', italic=True, color=GRAY, size=10)

add_heading(doc, 'Running Continuous Monitoring', level=2)
add_code(doc, 'python main.py monitor start --interval 60 --alert-email you@gmail.com')
add_para(doc, 'This runs forever in the terminal window, checking every 60 minutes. Press Ctrl+C to stop.', size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 11 — AI COPY
# ════════════════════════════════════════════════════════════════
add_heading(doc, '11. Module 5 - AI Ad Copy Generator', level=1)
add_divider(doc)
add_para(doc, 'This optional feature uses Claude AI (by Anthropic) to write professional Facebook and Instagram ad copy based on your product and audience.', size=11)
doc.add_paragraph()

add_heading(doc, 'What It Generates', level=2)
add_table(doc,
    ['Field', 'Description', 'Character Limit'],
    [
        ('Headline',          'Main attention-grabbing title',                   '40 characters'),
        ('Body',              'The main ad text (Pain -> Solution -> Proof)',    '125 characters'),
        ('CTA',               'Call-to-action button text',                      'Fixed options'),
        ('Hook',              'First 3 words optimized to stop scrolling',       '~15 characters'),
        ('Compliance Note',   'Flags any claims that might need substantiation', 'N/A'),
    ]
)

add_heading(doc, 'Psychological Angles Used', level=2)
add_bullet(doc, 'Variation 1: Curiosity - makes people want to know more')
add_bullet(doc, 'Variation 2: Social Proof - others are doing it, you should too')
add_bullet(doc, 'Variation 3: Urgency / FOMO - limited time, act now')

add_heading(doc, 'Example Command', level=2)
add_code(doc, 'python main.py copy generate \\')
add_code(doc, '  --product "Running shoes for women" \\')
add_code(doc, '  --audience "Fitness-conscious women aged 25-35 in India" \\')
add_code(doc, '  --tone urgent \\')
add_code(doc, '  --objective sales \\')
add_code(doc, '  --count 3')

add_heading(doc, 'Available Tones', level=2)
add_table(doc,
    ['Tone', 'Best For'],
    [
        ('professional', 'B2B, finance, healthcare'),
        ('casual',       'Lifestyle, food, entertainment'),
        ('urgent',       'Sales, limited offers, flash deals'),
        ('emotional',    'Charity, personal care, family products'),
        ('humorous',     'Youth brands, fun products'),
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 12 — DATABASE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '12. Module 6 - Local Database', level=1)
add_divider(doc)
add_para(doc, 'The system stores everything in a local SQLite database file (meta_ads.db). This means reports and stats are available instantly without making API calls.', size=11)
doc.add_paragraph()

add_table(doc,
    ['Table', 'What Is Stored', 'Key Columns'],
    [
        ('campaigns',         'All your campaigns',               'id, name, objective, status, daily_budget'),
        ('adsets',            'All ad sets',                      'id, campaign_id, name, status, daily_budget, targeting'),
        ('ads',               'All individual ads',               'id, adset_id, name, headline, body, cta, image_url'),
        ('insights',          'Daily performance metrics',        'entity_id, date, impressions, clicks, spend, roas, ctr'),
        ('generated_copy',    'AI-written ad copy history',       'headline, body, cta, product, audience, tone'),
        ('alerts',            'All alerts ever triggered',        'alert_type, campaign_id, actual_value, threshold'),
        ('optimization_log',  'Every budget change ever made',    'entity_id, action, old_value, new_value, reason, dry_run'),
    ]
)

add_heading(doc, 'Key Design Decisions', level=2)
add_bullet(doc, 'Budgets stored in cents (not rupees) to match Meta API format exactly')
add_bullet(doc, 'Insights table uses UNIQUE constraint so re-syncing never creates duplicates')
add_bullet(doc, 'optimization_log tracks dry_run=1 for preview runs vs dry_run=0 for real changes')
add_bullet(doc, 'Database file is created automatically on first run - no setup needed')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 13 — CONFIGURATION
# ════════════════════════════════════════════════════════════════
add_heading(doc, '13. Configuration & Thresholds', level=1)
add_divider(doc)
add_para(doc, 'All thresholds are in config.py and can also be overridden in your .env file.', size=11)
doc.add_paragraph()

add_table(doc,
    ['Setting', 'Default Value', 'What It Controls'],
    [
        ('MIN_ROAS_THRESHOLD',          '1.5',    'Pause ad set if ROAS falls below this'),
        ('SCALE_ROAS_THRESHOLD',        '3.0',    'Increase budget if ROAS exceeds this'),
        ('MAX_CPA_THRESHOLD',           '25.0',   'Decrease budget if cost/lead exceeds this (USD)'),
        ('SCALE_CPA_THRESHOLD',         '15.0',   'Increase budget if cost/lead is below this (USD)'),
        ('MIN_CTR_THRESHOLD',           '0.5%',   'Alert if CTR falls below this'),
        ('MAX_FREQUENCY',               '4.0',    'Alert if average frequency exceeds this'),
        ('SPEND_SPIKE_MULTIPLIER',      '2.0',    'Alert if today spend is 2x the daily average'),
        ('MIN_SPEND_FOR_ROAS_EVAL',     '$50',    'Minimum spend before evaluating ROAS'),
        ('BUDGET_INCREASE_PCT',         '15%',    'How much to increase budget for winners'),
        ('BUDGET_DECREASE_PCT',         '20%',    'How much to cut budget for losers'),
        ('MAX_ADSET_BUDGET_MULTIPLIER', '5x',     'Maximum budget is 5x the original'),
        ('MIN_ADSET_BUDGET_CENTS',      '500',    'Budget never goes below Rs.5 (or $5)'),
    ]
)

add_para(doc, 'To override any threshold for your account, add it to your .env file:', size=11, indent=False)
add_code(doc, 'MIN_ROAS_THRESHOLD=2.0')
add_code(doc, 'MAX_CPA_THRESHOLD=1500')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 14 — WORKFLOW
# ════════════════════════════════════════════════════════════════
add_heading(doc, '14. Full Day-to-Day Workflow', level=1)
add_divider(doc)

add_heading(doc, 'Day 1 - Launch a New Campaign', level=2)
add_code(doc, '# 1. Generate AI copy for your ad')
add_code(doc, 'python main.py copy generate --product "Your product" --audience "Your audience" --tone urgent --count 3')
add_code(doc, '')
add_code(doc, '# 2. Create the campaign (dry run first to preview)')
add_code(doc, 'python main.py campaign create --name "My Campaign" --objective OUTCOME_TRAFFIC --budget 500 --dry-run')
add_code(doc, '')
add_code(doc, '# 3. Create it for real')
add_code(doc, 'python main.py campaign create --name "My Campaign" --objective OUTCOME_TRAFFIC --budget 500')
add_code(doc, '')
add_code(doc, '# 4. Create an ad in the ad set (use the Ad Set ID from step 3)')
add_code(doc, 'python main.py ad create --adset-id ADSET_ID --campaign-id CAMPAIGN_ID --name "Ad 1" --headline "Your headline" --body "Your body text" --cta SHOP_NOW --link https://yoursite.com')
add_code(doc, '')
add_code(doc, '# 5. Go to Meta Ads Manager and activate the campaign manually')

add_heading(doc, 'Every Day - Monitoring & Optimization', level=2)
add_code(doc, '# Check for alerts (quick check)')
add_code(doc, 'python main.py monitor check')
add_code(doc, '')
add_code(doc, '# Run optimization (previews first)')
add_code(doc, 'python main.py optimize --dry-run')
add_code(doc, '')
add_code(doc, '# Apply optimization if you are happy with the suggested changes')
add_code(doc, 'python main.py optimize')

add_heading(doc, 'Weekly - Performance Review', level=2)
add_code(doc, '# Campaign report for last 7 days')
add_code(doc, 'python main.py report --campaign-id YOUR_ID --date-range last_7d')
add_code(doc, '')
add_code(doc, '# Full account report')
add_code(doc, 'python main.py report account --date-range last_30d')
add_code(doc, '')
add_code(doc, '# Export to CSV for your records')
add_code(doc, 'python main.py report --campaign-id YOUR_ID --format csv --output weekly_report.csv')

add_heading(doc, 'Continuous (24/7) Monitoring Mode', level=2)
add_para(doc, 'Leave this running in a terminal window and the tool will check every 60 minutes automatically:', size=11)
add_code(doc, 'python main.py monitor start --interval 60 --alert-email your@email.com')
add_para(doc, 'Press Ctrl+C to stop monitoring.', italic=True, color=GRAY, size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 15 — TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════
add_heading(doc, '15. Troubleshooting', level=1)
add_divider(doc)

add_table(doc,
    ['Error', 'Cause', 'Fix'],
    [
        ('Missing required env vars',        '.env file not filled in or missing',           'Run: copy .env.example .env  then fill in credentials'),
        ('ModuleNotFoundError',              'Dependencies not installed',                   'Run: pip install -r requirements.txt'),
        ('Invalid access token',             'Token expired (short-lived lasts 1 hour)',     'Generate new token at developers.facebook.com/tools/explorer'),
        ('No campaigns found',               'New account with no campaigns yet',            'Create a campaign first or run --sync to sync from Meta'),
        ('No insights for adset, skipping',  'Ad set has no spend data yet',                'Normal for new campaigns - run after campaign has some spend'),
        ('API permission error',             'Access token missing required permissions',    'Regenerate token with: ads_management, ads_read, business_management'),
        ('UnicodeEncodeError',               'Terminal encoding issue on Windows',           'Run: chcp 65001 in terminal to enable UTF-8'),
        ('Email alert not sending',          'SMTP credentials not set or wrong',            'Check SMTP_USER, SMTP_PASS in .env - use Gmail App Password'),
    ]
)

doc.add_paragraph()
add_heading(doc, 'Check the Log File', level=2)
add_para(doc, 'All activity is logged to automation.log in the project folder. Check it when something goes wrong:', size=11)
add_code(doc, 'notepad automation.log')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  FINAL PAGE
# ════════════════════════════════════════════════════════════════
add_heading(doc, 'Quick Reference Card', level=1)
add_divider(doc)
add_table(doc,
    ['Task', 'Command'],
    [
        ('Test connection',              'python main.py setup verify'),
        ('Sync campaigns from Meta',     'python main.py campaign list --sync'),
        ('Create campaign (preview)',    'python main.py campaign create --name "X" --objective OUTCOME_TRAFFIC --budget 500 --dry-run'),
        ('Create campaign (real)',       'python main.py campaign create --name "X" --objective OUTCOME_TRAFFIC --budget 500'),
        ('Generate AI ad copy',          'python main.py copy generate --product "X" --audience "Y" --tone urgent --count 3'),
        ('Run optimization (preview)',   'python main.py optimize --dry-run'),
        ('Run optimization (real)',      'python main.py optimize'),
        ('View campaign report',         'python main.py report --campaign-id ID --date-range last_7d'),
        ('Export report to CSV',         'python main.py report --campaign-id ID --format csv --output file.csv'),
        ('Check alerts now',             'python main.py monitor check'),
        ('Start 24/7 monitoring',        'python main.py monitor start --interval 60'),
        ('View database stats',          'python main.py stats'),
        ('Pause a campaign',             'python main.py campaign pause --id CAMPAIGN_ID'),
        ('Resume a campaign',            'python main.py campaign resume --id CAMPAIGN_ID'),
    ],
    header_color='1877F2'
)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Meta Ads Automation v1.0  |  Built with Meta Marketing API + Python  |  2026')
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.italic = True

# ── Save ──────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'Meta Ads Automation - System Documentation.docx')
doc.save(output_path)
print(f"Document saved: {output_path}")
